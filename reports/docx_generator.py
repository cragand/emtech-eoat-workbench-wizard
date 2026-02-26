"""DOCX report generator for QC and maintenance workflows."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os


class DOCXReportGenerator:
    """Generate DOCX reports for QC and maintenance sessions."""
    
    def __init__(self, output_dir="output/reports"):
        """Initialize DOCX generator.
        
        Args:
            output_dir: Directory to save generated reports
        """
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
    
    def generate_report(self, serial_number, technician, description, images, mode_name="General Capture", 
                       workflow_name=None, checklist_data=None, video_paths=None, barcode_scans=None):
        """Generate a DOCX report.
        
        Args:
            serial_number: Serial number or identifier
            technician: Technician name
            description: Job description
            images: List of image file paths OR list of dicts with {path, camera, notes, barcode_scans}
            mode_name: Name of the mode used
            workflow_name: Optional workflow name (for Mode 2/3)
            checklist_data: Optional list of checklist items with status
            video_paths: Optional list of video file paths
            barcode_scans: Optional list of barcode scans {type, data, timestamp}
            
        Returns:
            Path to generated DOCX file
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        serial = serial_number if serial_number else "unknown"
        filename = f"{serial}_{timestamp}.docx"
        filepath = os.path.join(self.output_dir, filename)
        
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Title - determine report type based on mode_name
        if "Mode 1" in mode_name or "General" in mode_name or "Capture" in mode_name:
            report_title = "Emtech EOAT Report - Inspection"
        elif "Mode 2" in mode_name or "QC" in mode_name:
            report_title = "Emtech EOAT Report - QC"
        elif "Mode 3" in mode_name or "Maintenance" in mode_name:
            report_title = "Emtech EOAT Report - Maintenance/Repair"
        else:
            report_title = "Emtech EOAT Report - Inspection"
        
        title = doc.add_heading(report_title, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.color.rgb = RGBColor(119, 194, 94)  # Emtech green #77C25E
        
        # Session Information
        doc.add_heading('Session Information', level=2)
        
        # Calculate number of rows needed
        num_rows = 5
        if workflow_name:
            num_rows += 1
        if barcode_scans:
            num_rows += 1
        
        # Info table
        table = doc.add_table(rows=num_rows, cols=2)
        table.style = 'Light Grid Accent 1'
        
        row_idx = 0
        
        # Serial Number
        table.rows[row_idx].cells[0].text = 'Serial Number:'
        table.rows[row_idx].cells[1].text = serial_number if serial_number else "N/A"
        row_idx += 1
        
        # Technician
        table.rows[row_idx].cells[0].text = 'Technician:'
        table.rows[row_idx].cells[1].text = technician if technician else "N/A"
        row_idx += 1
        
        # Mode
        table.rows[row_idx].cells[0].text = 'Mode:'
        table.rows[row_idx].cells[1].text = mode_name
        row_idx += 1
        
        # Date/Time
        table.rows[row_idx].cells[0].text = 'Date/Time:'
        table.rows[row_idx].cells[1].text = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        row_idx += 1
        
        # Workflow (if applicable)
        if workflow_name:
            table.rows[row_idx].cells[0].text = 'Workflow:'
            table.rows[row_idx].cells[1].text = workflow_name
            row_idx += 1
        
        # Barcode Scans (if any) - show unique scans
        if barcode_scans:
            # Get unique barcode data (type + data combination)
            unique_scans = {}
            for scan in barcode_scans:
                key = f"{scan.get('type', 'Unknown')}: {scan.get('data', '')}"
                unique_scans[key] = True
            
            table.rows[row_idx].cells[0].text = 'Scan Info:'
            table.rows[row_idx].cells[1].text = "\n".join(unique_scans.keys())
            row_idx += 1
        
        # Make first column bold - safely
        for row in table.rows:
            cell = row.cells[0]
            # Get the text, clear it, and re-add with bold
            text = cell.text
            cell.text = ''
            run = cell.paragraphs[0].add_run(text)
            run.font.bold = True
        
        doc.add_paragraph()
        
        # Description as separate section
        if description:
            doc.add_heading('Description', level=3)
            doc.add_paragraph(description)
            doc.add_paragraph()
        
        # Procedure Summary (if provided)
        if checklist_data:
            doc.add_heading('Procedure Summary', level=2)
            
            summary_table = doc.add_table(rows=len(checklist_data) + 1, cols=2)
            summary_table.style = 'Light Grid Accent 1'
            
            # Header row
            summary_table.rows[0].cells[0].text = 'Step'
            summary_table.rows[0].cells[1].text = 'Status'
            for cell in summary_table.rows[0].cells:
                cell.paragraphs[0].runs[0].font.bold = True
            
            # Data rows
            for idx, item in enumerate(checklist_data, 1):
                summary_table.rows[idx].cells[0].text = item['name']
                
                if item.get('has_pass_fail', False):
                    status = "✓ Pass" if item.get('passed', False) else "✗ Fail"
                    status_cell = summary_table.rows[idx].cells[1]
                    status_cell.text = status
                    if item.get('passed', False):
                        status_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(76, 175, 80)  # Green
                    else:
                        status_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(244, 67, 54)  # Red
                else:
                    status = "✓ Complete"
                    status_cell = summary_table.rows[idx].cells[1]
                    status_cell.text = status
                    status_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(129, 199, 132)  # Light green
            
            doc.add_paragraph()
        
        # Checklist (if provided)
        if checklist_data:
            doc.add_heading('Procedure Steps', level=2)
            
            for item in checklist_data:
                # Step name
                p = doc.add_paragraph()
                p.add_run(item['name']).bold = True
                
                # Status - Complete/Incomplete or Pass/Fail
                if item.get('has_pass_fail', False):
                    status = "✓ Pass" if item.get('passed', False) else "✗ Fail"
                    status_run = p.add_run(f" - {status}")
                    status_run.bold = True
                    if item.get('passed', False):
                        status_run.font.color.rgb = RGBColor(76, 175, 80)  # Green
                    else:
                        status_run.font.color.rgb = RGBColor(244, 67, 54)  # Red
                else:
                    status = "✓ Complete"
                    status_run = p.add_run(f" - {status}")
                    status_run.bold = True
                    status_run.font.color.rgb = RGBColor(129, 199, 132)  # Light green
                
                # Description
                if item.get('description'):
                    desc_text = item['description'][:200] + "..." if len(item['description']) > 200 else item['description']
                    p = doc.add_paragraph()
                    p.add_run(desc_text).italic = True
                
                # Reference image with label
                if item.get('checkbox_image') and os.path.exists(item['checkbox_image']):
                    p = doc.add_paragraph()
                    p.add_run('Reference Image (Inspection Points):').bold = True
                    try:
                        doc.add_picture(item['checkbox_image'], width=Inches(4))
                    except Exception as e:
                        p = doc.add_paragraph()
                        p.add_run(f'Error loading reference image: {str(e)}').italic = True
                
                # Captured images for this step
                step_number = item.get('step_number')
                if step_number and images:
                    step_images = [img for img in images if isinstance(img, dict) and img.get('step') == step_number]
                    
                    if step_images:
                        p = doc.add_paragraph()
                        p.add_run(f'Captured Images ({len(step_images)}):').bold = True
                        
                        for img_data in step_images:
                            img_path = img_data['path']
                            camera = img_data.get('camera', 'Unknown')
                            notes = img_data.get('notes', '')
                            markers = img_data.get('markers', [])
                            
                            if os.path.exists(img_path):
                                p = doc.add_paragraph()
                                p.add_run('Camera: ').italic = True
                                p.add_run(camera)
                                
                                if notes:
                                    p = doc.add_paragraph()
                                    p.add_run('Notes: ').bold = True
                                    p.add_run(notes)
                                
                                marker_notes = [m for m in markers if m.get('note', '').strip()]
                                if marker_notes:
                                    p = doc.add_paragraph()
                                    p.add_run('Annotations:').bold = True
                                    for m in marker_notes:
                                        p = doc.add_paragraph(style='List Bullet')
                                        p.add_run(f"{m['label']}: {m['note']}")
                                
                                try:
                                    doc.add_picture(img_path, width=Inches(4))
                                except Exception as e:
                                    p = doc.add_paragraph()
                                    p.add_run(f'Error loading image: {str(e)}').italic = True
                                
                                doc.add_paragraph()
                
                doc.add_paragraph()  # Spacing between steps
            
            doc.add_paragraph()
        
        # Captured Images (only for Mode 1 - no workflow)
        elif images:
            doc.add_heading(f'Captured Images ({len(images)})', level=2)
            
            for idx, img_data in enumerate(images, 1):
                # Handle both old format (string path) and new format (dict with metadata)
                if isinstance(img_data, dict):
                    img_path = img_data['path']
                    camera = img_data.get('camera', 'Unknown')
                    notes = img_data.get('notes', '')
                    media_type = img_data.get('type', 'image')
                    step_info = img_data.get('step_info', '')
                    markers = img_data.get('markers', [])
                    img_barcode_scans = img_data.get('barcode_scans', [])
                else:
                    # Legacy format - just a path string
                    img_path = img_data
                    camera = 'Unknown'
                    notes = ''
                    media_type = 'image'
                    step_info = ''
                    markers = []
                    img_barcode_scans = []
                
                if os.path.exists(img_path):
                    # Check if this is a video file
                    is_video = media_type == 'video' or img_path.lower().endswith(('.avi', '.mp4', '.mov', '.mkv'))
                    
                    if is_video:
                        # For videos, just add a note
                        p = doc.add_paragraph()
                        p.add_run(f'Video {idx}: ').bold = True
                        p.add_run(os.path.basename(img_path))
                        
                        p = doc.add_paragraph()
                        p.add_run('Camera: ').italic = True
                        p.add_run(camera)
                        
                        p = doc.add_paragraph()
                        p.add_run('(Video file - not embedded in report)').italic = True
                        
                        if notes:
                            p = doc.add_paragraph()
                            p.add_run('Notes:').bold = True
                            # Handle multi-line notes by splitting on newlines
                            for line in notes.split('\n'):
                                p = doc.add_paragraph(line)
                                p.paragraph_format.left_indent = Inches(0.25)
                        
                        # Add marker notes if present
                        marker_notes = [m for m in markers if m.get('note', '').strip()]
                        if marker_notes:
                            p = doc.add_paragraph()
                            p.add_run('Annotations:').bold = True
                            for m in marker_notes:
                                p = doc.add_paragraph(style='List Bullet')
                                p.add_run(f"{m['label']}: {m['note']}")
                        
                        if step_info:
                            p = doc.add_paragraph()
                            p.add_run('Step: ').italic = True
                            p.add_run(step_info)
                        
                        doc.add_paragraph()
                    else:
                        # For images, embed them
                        p = doc.add_paragraph()
                        p.add_run(f'Image {idx}: ').bold = True
                        p.add_run(os.path.basename(img_path))
                        
                        p = doc.add_paragraph()
                        p.add_run('Camera: ').italic = True
                        p.add_run(camera)
                        
                        if notes:
                            p = doc.add_paragraph()
                            p.add_run('Notes:').bold = True
                            # Handle multi-line notes by splitting on newlines
                            for line in notes.split('\n'):
                                p = doc.add_paragraph(line)
                                p.paragraph_format.left_indent = Inches(0.25)
                        
                        # Add marker notes if present
                        marker_notes = [m for m in markers if m.get('note', '').strip()]
                        if marker_notes:
                            p = doc.add_paragraph()
                            p.add_run('Annotations:').bold = True
                            for m in marker_notes:
                                p = doc.add_paragraph(style='List Bullet')
                                p.add_run(f"{m['label']}: {m['note']}")
                        
                        
                        if step_info:
                            p = doc.add_paragraph()
                            p.add_run('Step: ').italic = True
                            p.add_run(step_info)
                        
                        try:
                            # Add image (scaled to fit page width - 6 inches)
                            doc.add_picture(img_path, width=Inches(6))
                        except Exception as e:
                            # If image can't be loaded, show error message
                            p = doc.add_paragraph()
                            p.add_run(f'Error loading image: {str(e)}').italic = True
                        
                        doc.add_paragraph()
        else:
            doc.add_paragraph('No images captured')
        
        # Recorded Videos section
        if video_paths:
            doc.add_page_break()
            doc.add_heading(f'Recorded Videos ({len(video_paths)})', level=2)
            
            for idx, video_path in enumerate(video_paths, 1):
                # Get relative path from output directory
                rel_path = os.path.relpath(video_path, start=os.path.dirname(filepath))
                p = doc.add_paragraph(style='List Number')
                p.add_run(rel_path).bold = True
            
            doc.add_paragraph()
            note = doc.add_paragraph()
            note.add_run('Note: Video files are saved separately and not embedded in this report.').italic = True
        
        # Save document
        doc.save(filepath)
        return filepath


def create_simple_docx_report(serial_number, description, images, output_dir="output/reports"):
    """Convenience function to quickly generate a simple DOCX report.
    
    Args:
        serial_number: Serial number or identifier
        description: Job description
        images: List of image file paths OR list of dicts with {path, camera, notes}
        output_dir: Output directory for report
        
    Returns:
        Path to generated DOCX
    """
    generator = DOCXReportGenerator(output_dir)
    return generator.generate_report(serial_number, description, images)
